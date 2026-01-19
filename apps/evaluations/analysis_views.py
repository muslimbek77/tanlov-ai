"""
Tender Tahlil API

Tender shartnomasi va ishtirokchilarni tahlil qilish uchun API endpointlar.
"""
from rest_framework import viewsets, status
from rest_framework.decorators import api_view, action, permission_classes
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from rest_framework.permissions import AllowAny
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.http import HttpResponse
import os
import logging
import PyPDF2
from docx import Document
import tempfile
import io
from datetime import datetime

# PDF uchun
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import registerFontFamily

from core.tender_analyzer import tender_analyzer
from core.services import document_processor

logger = logging.getLogger(__name__)


def _normalize_language(raw_language: str) -> str:
    lang = (raw_language or 'uz_latn').strip().lower()
    if lang in ['uz', 'uz-latn', 'uz_uz']:
        return 'uz_latn'
    if lang in ['uz-cyrl', 'uz_cyrl', 'uz-cyrl-uz']:
        return 'uz_cyrl'
    if lang in ['ru', 'ru-ru', 'rus']:
        return 'ru'
    return lang


def _is_uzbek(lang: str) -> bool:
    return _normalize_language(lang).startswith('uz')


def _register_pdf_fonts() -> dict:
    candidates = [
        (
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
            'DejaVuSans',
        ),
        (
            '/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf',
            'DejaVuSerif',
        ),
    ]
    try:
        import reportlab
        base = os.path.join(os.path.dirname(reportlab.__file__), 'fonts')
        candidates.append((os.path.join(base, 'Vera.ttf'), os.path.join(base, 'VeraBd.ttf'), 'Vera'))
    except Exception:
        pass

    for regular_path, bold_path, family in candidates:
        if os.path.exists(regular_path) and os.path.exists(bold_path):
            try:
                pdfmetrics.registerFont(TTFont(family, regular_path))
                pdfmetrics.registerFont(TTFont(f'{family}-Bold', bold_path))
                registerFontFamily(family, normal=family, bold=f'{family}-Bold', italic=family, boldItalic=f'{family}-Bold')
                return {'family': family, 'regular': family, 'bold': f'{family}-Bold'}
            except Exception:
                continue

    return {'family': 'Helvetica', 'regular': 'Helvetica', 'bold': 'Helvetica-Bold'}


def _msg(language: str, uz_latn: str, uz_cyrl: str, ru: str) -> str:
    lang = _normalize_language(language)
    if lang == 'ru':
        return ru
    if lang == 'uz_cyrl':
        return uz_cyrl
    return uz_latn


def extract_text_from_file(file) -> str:
    """Fayldan matn ajratib olish"""
    text = ""
    filename = file.name.lower()
    
    try:
        # Vaqtinchalik faylga saqlash
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp:
            for chunk in file.chunks():
                tmp.write(chunk)
            tmp_path = tmp.name
        
        if filename.endswith('.pdf'):
            with open(tmp_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        
        elif filename.endswith('.docx'):
            doc = Document(tmp_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        
        elif filename.endswith('.txt'):
            with open(tmp_path, 'r', encoding='utf-8') as f:
                text = f.read()
        
        else:
            # Oddiy matn sifatida o'qish
            text = file.read().decode('utf-8', errors='ignore')
        
        # Vaqtinchalik faylni o'chirish
        os.unlink(tmp_path)
        
    except Exception as e:
        logger.error(f"Fayl o'qishda xatolik: {str(e)}")
        text = ""
    
    return text


@api_view(['POST'])
@permission_classes([AllowAny])
def analyze_tender(request):
    """
    Tender shartnomasini tahlil qilish
    
    POST /api/evaluations/analyze-tender/
    
    Body:
        - file: Tender shartnoma fayli (PDF, DOCX, TXT)
        - text: Yoki to'g'ridan-to'g'ri matn
        - metadata: Qo'shimcha ma'lumotlar (JSON)
        - language: Til (uz yoki ru)
    
    Returns:
        - success: bool
        - analysis: Tender tahlili natijalari
    """
    try:
        tender_text = ""
        metadata = {}
        language = _normalize_language(request.data.get('language', 'uz_latn'))
        
        # Fayldan yoki matndan olish
        if 'file' in request.FILES:
            file = request.FILES['file']
            allowed_ext = ['.pdf', '.docx', '.txt']
            ext = os.path.splitext(file.name)[1].lower()
            if ext not in allowed_ext:
                error_msg = _msg(
                    language,
                    'Faqat PDF, DOCX yoki TXT fayl yuklang',
                    'Фақат PDF, DOCX ёки TXT файл юкланг',
                    'Загрузите только PDF, DOCX или TXT файл',
                )
                return Response({
                    'success': False,
                    'error': error_msg
                }, status=status.HTTP_400_BAD_REQUEST)
            tender_text = extract_text_from_file(file)
            metadata['filename'] = file.name
            metadata['file_size'] = file.size
        # Mazmun va mantiqiy tekshiruv (GPT orqali)
        if tender_text.strip():
            from core.llm_engine import llm_engine
            if not getattr(llm_engine, 'providers', None):
                error_msg = _msg(
                    language,
                    "LLM provayderi sozlanmagan. .env faylga OPENAI_API_KEY qo'shing yoki Ollama'ni ishga tushiring.",
                    "LLM провайдери созланмаган. .env файлга OPENAI_API_KEY қўшинг ёки Ollama'ни ишга туширинг.",
                    "LLM провайдер не настроен. Добавьте OPENAI_API_KEY в .env или запустите Ollama.",
                )
                return Response({
                    'success': False,
                    'error': error_msg
                }, status=status.HTTP_503_SERVICE_UNAVAILABLE)

            check_prompt = f"""
Quyidagi matn tender shartnomasi yoki unga o'xshash rasmiy hujjatmi? Mazmunan va mantiqan to'g'ri, to'liq va tahlil qilishga yaroqlimi? Faqat 'HA' yoki 'YO' deb javob ber:

{tender_text[:2000]}
"""
            check_result = llm_engine.generate_response(check_prompt, system_prompt="Faqat 'HA' yoki 'YO' deb javob ber.", temperature=0.0)
            if isinstance(check_result, dict):
                check_response = check_result.get('response', '').strip().upper()
            else:
                check_response = str(check_result).strip().upper()
            if 'YO' in check_response:
                error_msg = _msg(
                    language,
                    "Yuklangan fayl mazmunan yoki mantiqan noto'g'ri. Iltimos, to'g'ri tender faylini yuklang.",
                    "Юкланган файл мазмунан ёки мантиқан нотўғри. Илтимос, тўғри тендер файлини юкланг.",
                    "Загруженный файл не соответствует содержанию тендера. Пожалуйста, загрузите правильный файл.",
                )
                return Response({
                    'success': False,
                    'error': error_msg
                }, status=status.HTTP_400_BAD_REQUEST)
        elif 'text' in request.data:
            tender_text = request.data.get('text', '')
        else:
            error_msg = _msg(
                language,
                'Tender fayli yoki matni talab qilinadi',
                'Тендер файли ёки матни талаб қилинади',
                'Требуется файл или текст тендера',
            )
            return Response({
                'success': False,
                'error': error_msg
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if not tender_text.strip():
            error_msg = _msg(
                language,
                "Tender matni bo'sh",
                "Тендер матни бўш",
                'Текст тендера пуст',
            )
            return Response({
                'success': False,
                'error': error_msg
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Qo'shimcha metadata
        if 'metadata' in request.data:
            try:
                import json
                extra_meta = json.loads(request.data.get('metadata', '{}'))
                metadata.update(extra_meta)
            except:
                pass
        
        metadata['language'] = language
        
        # Tahlil qilish
        result = tender_analyzer.analyze_tender_document(tender_text, metadata)
        
        if result['success']:
            return Response(result, status=status.HTTP_200_OK)

        error_text = str(result.get('error', ''))
        if 'Hech qanday LLM provayderi mavjud emas' in error_text or 'LLM' in error_text:
            error_msg = _msg(
                language,
                "LLM provayderi sozlanmagan yoki mavjud emas. .env faylga OPENAI_API_KEY qo'shing yoki Ollama'ni sozlang.",
                "LLM провайдери созланмаган ёки мавжуд эмас. .env файлга OPENAI_API_KEY қўшинг ёки Ollama'ни созланг.",
                "LLM провайдер не настроен или недоступен. Добавьте OPENAI_API_KEY в .env или настройте Ollama.",
            )
            return Response({
                'success': False,
                'error': error_msg
            }, status=status.HTTP_503_SERVICE_UNAVAILABLE)

        return Response(result, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
    except Exception as e:
        logger.error(f"Tender tahlilida xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([AllowAny])
def analyze_participant(request):
    """
    Ishtirokchi hujjatlarini tahlil qilish
    
    POST /api/evaluations/analyze-participant/
    
    Body:
        - name: Ishtirokchi nomi
        - file: Ishtirokchi hujjati (PDF, DOCX, TXT)
        - text: Yoki to'g'ridan-to'g'ri matn
        - metadata: Qo'shimcha ma'lumotlar (JSON)
        - language: Til (uz yoki ru)
        - tender_data: Tender tahlili (agar server xotirasida yo'q bo'lsa)
    
    Returns:
        - success: bool
        - analysis: Ishtirokchi tahlili natijalari
    """
    try:
        language = _normalize_language(request.data.get('language', 'uz_latn'))
        default_name = _msg(language, "Noma'lum ishtirokchi", "Номаълум иштирокчи", 'Неизвестный участник')
        participant_name = request.data.get('name', default_name)
        participant_text = ""
        metadata = {'language': language}
        
        logger.info(f"Ishtirokchi tahlili so'rovi: {participant_name}")
        logger.info(f"tender_requirements mavjud: {len(tender_analyzer.tender_requirements)}")
        
        # Agar tender tahlili yo'q bo'lsa, frontend'dan kelgan ma'lumotni tiklash
        if not tender_analyzer.tender_requirements:
            tender_data_str = request.data.get('tender_data', '')
            logger.info(f"tender_data mavjud: {bool(tender_data_str)}, uzunligi: {len(str(tender_data_str)) if tender_data_str else 0}")
            
            if tender_data_str:
                try:
                    import json
                    tender_data = json.loads(tender_data_str) if isinstance(tender_data_str, str) else tender_data_str
                    logger.info(f"tender_data parse qilindi, kalit: {list(tender_data.keys()) if isinstance(tender_data, dict) else 'dict emas'}")
                    # Tender ma'lumotlarini tiklash
                    restored = tender_analyzer.restore_tender_analysis(tender_data)
                    logger.info(f"Tender tahlili tiklandi: {restored}, requirements: {len(tender_analyzer.tender_requirements)}")
                except Exception as e:
                    logger.warning(f"Tender ma'lumotlarini tiklashda xatolik: {e}")
                    import traceback
                    traceback.print_exc()
        
        # Fayldan yoki matndan olish
        if 'file' in request.FILES:
            file = request.FILES['file']
            allowed_ext = ['.pdf', '.docx', '.txt']
            ext = os.path.splitext(file.name)[1].lower()
            if ext not in allowed_ext:
                error_msg = _msg(
                    language,
                    'Faqat PDF, DOCX yoki TXT fayl yuklang',
                    'Фақат PDF, DOCX ёки TXT файл юкланг',
                    'Загрузите только PDF, DOCX или TXT файл',
                )
                return Response({
                    'success': False,
                    'error': error_msg
                }, status=status.HTTP_400_BAD_REQUEST)
            participant_text = extract_text_from_file(file)
            metadata['filename'] = file.name
            metadata['file_size'] = file.size
        # Mazmun va mantiqiy tekshiruv (GPT orqali)
        if participant_text.strip():
            from core.llm_engine import llm_engine
            check_prompt = f"""
Quyidagi matn ishtirokchi hujjatlari yoki unga o'xshash rasmiy hujjatmi? Mazmunan va mantiqan to'g'ri, to'liq va tahlil qilishga yaroqlimi? Faqat 'HA' yoki 'YO' deb javob ber:

{participant_text[:2000]}
"""
            check_result = llm_engine.generate_response(check_prompt, system_prompt="Faqat 'HA' yoki 'YO' deb javob ber.", temperature=0.0)
            if isinstance(check_result, dict):
                check_response = check_result.get('response', '').strip().upper()
            else:
                check_response = str(check_result).strip().upper()
            if 'YO' in check_response:
                error_msg = _msg(
                    language,
                    "Yuklangan fayl mazmunan yoki mantiqan noto'g'ri. Iltimos, to'g'ri ishtirokchi faylini yuklang.",
                    "Юкланган файл мазмунан ёки мантиқан нотўғри. Илтимос, тўғри иштирокчи файлини юкланг.",
                    "Загруженный файл не соответствует содержанию участника. Пожалуйста, загрузите правильный файл.",
                )
                return Response({
                    'success': False,
                    'error': error_msg
                }, status=status.HTTP_400_BAD_REQUEST)
        elif 'text' in request.data:
            participant_text = request.data.get('text', '')
        else:
            error_msg = _msg(
                language,
                'Ishtirokchi fayli yoki matni talab qilinadi',
                'Иштирокчи файли ёки матни талаб қилинади',
                'Требуется файл или текст участника',
            )
            return Response({
                'success': False,
                'error': error_msg
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if not participant_text.strip():
            error_msg = _msg(
                language,
                "Ishtirokchi matni bo'sh",
                "Иштирокчи матни бўш",
                'Текст участника пуст',
            )
            return Response({
                'success': False,
                'error': error_msg
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Qo'shimcha metadata
        if 'metadata' in request.data:
            try:
                import json
                extra_meta = json.loads(request.data.get('metadata', '{}'))
                metadata.update(extra_meta)
            except:
                pass
        
        # Tahlil qilish
        result = tender_analyzer.analyze_participant(
            participant_name, 
            participant_text, 
            metadata
        )
        
        if result['success']:
            return Response(result, status=status.HTTP_200_OK)
        else:
            return Response(result, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
    except Exception as e:
        logger.error(f"Ishtirokchi tahlilida xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([AllowAny])
def compare_participants(request):
    """
    Ishtirokchilarni solishtirish
    
    POST /api/evaluations/compare-participants/
    
    Body:
        - participants: Ishtirokchilar tahlillari ro'yxati
        - language: Til (uz yoki ru)
    
    Returns:
        - success: bool
        - ranking: Reyting bo'yicha saralangan ishtirokchilar
        - winner: G'olib
        - summary: Xulosa
    """
    try:
        participants = request.data.get('participants', [])
        language = _normalize_language(request.data.get('language', 'uz_latn'))
        
        if not participants:
            error_msg = _msg(
                language,
                "Ishtirokchilar ro'yxati bo'sh",
                "Иштирокчилар рўйхати бўш",
                'Список участников пуст',
            )
            return Response({
                'success': False,
                'error': error_msg
            }, status=status.HTTP_400_BAD_REQUEST)
        
        result = tender_analyzer.compare_participants(participants, language)
        
        if result['success']:
            return Response(result, status=status.HTTP_200_OK)
        else:
            return Response(result, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
    except Exception as e:
        logger.error(f"Ishtirokchilar solishtirishda xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([AllowAny])
def full_analysis(request):
    """
    To'liq tahlil - tender + barcha ishtirokchilar
    
    POST /api/evaluations/full-analysis/
    
    Body:
        - tender_file: Tender shartnoma fayli
        - tender_text: Yoki tender matni
        - participants: [
            {name: "Kompaniya nomi", file: fayl yoki text: matn}
        ]
    
    Returns:
        - success: bool
        - tender_analysis: Tender tahlili
        - participants_analysis: Ishtirokchilar tahlili
        - ranking: Reyting
        - winner: G'olib
        - summary: Xulosa
    """
    try:
        results = {
            'success': True,
            'tender_analysis': None,
            'participants_analysis': [],
            'ranking': [],
            'winner': None,
            'summary': '',
            'progress': {
                'step': 'init',
                'current': 0,
                'total': 1,
                'status': 'Tender tahlili boshlanmoqda'
            }
        }
        
        # 1. Tender tahlili
        tender_text = ""
        if 'tender_file' in request.FILES:
            tender_text = extract_text_from_file(request.FILES['tender_file'])
        elif 'tender_text' in request.data:
            tender_text = request.data.get('tender_text', '')
        
        if not tender_text.strip():
            results['progress'] = {
                'step': 'tender_error',
                'current': 0,
                'total': 1,
                'status': 'Tender fayli yoki matni talab qilinmadi'
            }
            return Response({
                'success': False,
                'error': 'Tender fayli yoki matni talab qilinadi',
                'progress': results['progress']
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Har doim yangi tender tahlil obyektini yaratamiz
        from core.tender_analyzer import TenderAnalyzer
        global tender_analyzer
        tender_analyzer = TenderAnalyzer()
        results['progress'] = {
            'step': 'tender_analysis',
            'current': 1,
            'total': 1,
            'status': 'Tender tahlili yakunlandi'
        }
        tender_result = tender_analyzer.analyze_tender_document(tender_text)
        if not tender_result['success']:
            results['progress'] = {
                'step': 'tender_error',
                'current': 1,
                'total': 1,
                'status': 'Tender tahlilida xatolik'
            }
            return Response({**tender_result, 'progress': results['progress']}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        # To'liq tender tahlilini saqlaymiz (requirements va boshqa maydonlar bilan)
        full_tender_analysis = tender_result['analysis'] if 'analysis' in tender_result else tender_result
        results['tender_analysis'] = full_tender_analysis
        
        # 2. Ishtirokchilar tahlili
        participants_data = request.data.get('participants', [])
        participant_names = set()
        participant_counter = 1
        total_participants = len([k for k in request.FILES if k.startswith('participant_')]) + len(participants_data)
        current_participant = 0
        # Fayldan
        for key in request.FILES:
            if key.startswith('participant_'):
                idx = key.replace('participant_', '').replace('_file', '')
                name = request.data.get(f'participant_{idx}_name', '').strip()
                if not name or name in participant_names:
                    name = f'Ishtirokchi {participant_counter}'
                file = request.FILES[key]
                text = extract_text_from_file(file)
                if text.strip():
                    tender_analyzer.restore_tender_analysis(full_tender_analysis)
                    participant_result = tender_analyzer.analyze_participant(name, text)
                    if participant_result['success']:
                        analysis = participant_result['analysis']
                        # Ensure risk_level is always present
                        if 'risk_level' not in analysis:
                            risk = analysis.get('risk_assessment', {}).get('overall_risk', 'unknown')
                            analysis['risk_level'] = risk
                        results['participants_analysis'].append(analysis)
                        participant_names.add(name)
                        participant_counter += 1
                current_participant += 1
                results['progress'] = {
                    'step': 'participant',
                    'current': current_participant,
                    'total': total_participants,
                    'status': f'{name} tahlil qilindi'
                }
        # JSON dan
        for p in participants_data:
            name = p.get('name', '').strip()
            if not name or name in participant_names:
                name = f'Ishtirokchi {participant_counter}'
            text = p.get('text', '')
            if text.strip():
                tender_analyzer.restore_tender_analysis(full_tender_analysis)
                participant_result = tender_analyzer.analyze_participant(name, text)
                if participant_result['success']:
                    analysis = participant_result['analysis']
                    if 'risk_level' not in analysis:
                        risk = analysis.get('risk_assessment', {}).get('overall_risk', 'unknown')
                        analysis['risk_level'] = risk
                    results['participants_analysis'].append(analysis)
                    participant_names.add(name)
                    participant_counter += 1
            current_participant += 1
            results['progress'] = {
                'step': 'participant',
                'current': current_participant,
                'total': total_participants,
                'status': f'{name} tahlil qilindi'
            }
        
        # 3. Solishtirish
        if len(results['participants_analysis']) < 2:
            error_msg = 'Reyting yaratish uchun kamida 2 ta ishtirokchi kerak. Hozircha {} ta tahlil qilingan. Yana ishtirokchi qo\'shing.'.format(len(results['participants_analysis']))
            results['progress'] = {
                'step': 'compare_error',
                'current': current_participant,
                'total': total_participants,
                'status': error_msg
            }
            return Response({
                'success': False,
                'error': error_msg,
                'progress': results['progress']
            }, status=status.HTTP_400_BAD_REQUEST)
        compare_result = tender_analyzer.compare_participants(results['participants_analysis'])
        if compare_result['success']:
            results['ranking'] = compare_result['ranking']
            results['winner'] = compare_result['winner']
            results['summary'] = compare_result['summary']
            results['progress'] = {
                'step': 'done',
                'current': current_participant,
                'total': total_participants,
                'status': 'Tahlil yakunlandi'
            }
        else:
            results['progress'] = {
                'step': 'compare_error',
                'current': current_participant,
                'total': total_participants,
                'status': compare_result.get('error', 'Solishtirishda xatolik')
            }
        return Response(results, status=status.HTTP_200_OK)
        
    except Exception as e:
        logger.error(f"To'liq tahlilda xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def get_tender_requirements(request):
    """
    Joriy tender talablarini olish
    
    GET /api/evaluations/tender-requirements/
    """
    requirements = tender_analyzer.get_tender_requirements()
    tender_info = tender_analyzer.get_tender_info()
    
    return Response({
        'success': True,
        'requirements': requirements,
        'tender_info': tender_info
    })


@api_view(['POST'])
def reset_analysis(request):
    """
    Tahlilni qayta boshlash
    
    POST /api/evaluations/reset/
    """
    global tender_analyzer
    from core.tender_analyzer import TenderAnalyzer
    tender_analyzer = TenderAnalyzer()
    
    return Response({
        'success': True,
        'message': 'Tahlil qayta boshlandi'
    })


def draw_header_footer(canvas, doc, language='uz'):
    """PDF sahifa header va footer - sodda va professional"""
    canvas.saveState()
    
    page_width, page_height = A4
    margin = 1.5*cm
    
    # ===== SODDA RAMKA =====
    canvas.setStrokeColor(colors.Color(0.2, 0.45, 0.3))
    canvas.setLineWidth(1.5)
    canvas.rect(margin, margin, page_width - 2*margin, page_height - 2*margin)
    
    # ===== HEADER =====
    header_y = page_height - margin - 15
    
    # TANLOV AI
    canvas.setFillColor(colors.Color(0.2, 0.45, 0.3))
    fonts = _register_pdf_fonts()
    canvas.setFont(fonts['bold'], 12)
    canvas.drawString(margin + 15, header_y, "TANLOV AI")
    
    # Sana
    canvas.setFont(fonts['regular'], 9)
    canvas.setFillColor(colors.Color(0.4, 0.4, 0.4))
    date_str = datetime.now().strftime('%d.%m.%Y')
    canvas.drawRightString(page_width - margin - 15, header_y, date_str)
    
    # Header chiziq
    canvas.setStrokeColor(colors.Color(0.2, 0.45, 0.3))
    canvas.setLineWidth(0.5)
    canvas.line(margin + 10, header_y - 10, page_width - margin - 10, header_y - 10)
    
    # ===== FOOTER =====
    footer_y = margin + 15
    
    # Footer chiziq
    canvas.line(margin + 10, footer_y + 8, page_width - margin - 10, footer_y + 8)
    
    # Tashkilot nomi
    canvas.setFillColor(colors.Color(0.4, 0.4, 0.4))
    canvas.setFont(fonts['regular'], 7)
    lang = _normalize_language(language)
    if lang == 'ru':
        org_name = "Управление цифровизации и внедрения ИКТ"
    elif lang == 'uz_cyrl':
        org_name = "Рақамлаштириш ва АКТни жорий қилиш бошқармаси"
    else:
        org_name = "Raqamlashtirish va AKTni joriy qilish boshqarmasi"
    canvas.drawString(margin + 15, footer_y, org_name)
    
    # Sahifa raqami
    canvas.drawRightString(page_width - margin - 15, footer_y, f"{doc.page}")
    
    canvas.restoreState()


@api_view(['POST'])
@permission_classes([AllowAny])
def export_pdf(request):
    """
    Tahlil natijalarini PDF formatida eksport qilish - Minimalist dizayn
    """
    try:
        data = request.data
        language = _normalize_language(data.get('language', 'uz_latn'))
        tender_analysis = data.get('tender_analysis', {})
        ranking = data.get('ranking', [])
        winner = data.get('winner', {})
        summary = data.get('summary', '')

        fonts = _register_pdf_fonts()
        
        # PDF yaratish
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=A4,
            rightMargin=2.5*cm,
            leftMargin=2.5*cm,
            topMargin=70,
            bottomMargin=70
        )
        
        # Ranglar - Professional
        DARK = colors.Color(0.15, 0.15, 0.15)
        GRAY = colors.Color(0.4, 0.4, 0.4)
        LIGHT_GRAY = colors.Color(0.96, 0.96, 0.96)
        PRIMARY = colors.Color(0.2, 0.4, 0.3)  # To'q yashil
        ACCENT = colors.Color(0.25, 0.5, 0.35)  # Yashil accent
        LIGHT_GREEN = colors.Color(0.93, 0.97, 0.94)  # Och yashil fon
        BORDER = colors.Color(0.8, 0.85, 0.8)
        GOLD = colors.Color(0.85, 0.65, 0.15)  # G'olib uchun oltin rang
        
        # Stillar - bir xil shrift o'lchamlari
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'Title',
            fontSize=18,
            spaceAfter=8,
            spaceBefore=0,
            alignment=1,
            textColor=PRIMARY,
            fontName=fonts['bold'],
        )
        
        heading_style = ParagraphStyle(
            'Heading',
            fontSize=11,
            spaceAfter=6,
            spaceBefore=10,
            textColor=PRIMARY,
            fontName=fonts['bold'],
        )
        
        normal_style = ParagraphStyle(
            'Normal',
            fontSize=9,
            spaceAfter=3,
            leading=12,
            textColor=DARK,
            fontName=fonts['regular'],
        )
        
        small_style = ParagraphStyle(
            'Small',
            fontSize=8,
            spaceAfter=2,
            leading=10,
            textColor=GRAY,
            fontName=fonts['regular'],
        )
        
        # Elementlar
        elements = []
        
        # Tilga qarab matnlar
        if language == 'uz_latn':
            title = "Tender Tahlili Hisoboti"
            tender_info_title = "Tender ma'lumotlari"
            purpose_label = "Maqsad"
            type_label = "Tur"
            req_count_label = "Talablar"
            ranking_title = "Ishtirokchilar reytingi"
            winner_title = "G'olib"
            score_label = "Ball"
            match_label = "Moslik"
            risk_label = "Xavf"
            strengths_label = "Kuchli tomonlar"
            weaknesses_label = "Kamchiliklar"
            summary_title = "Xulosa"
            risk_levels = {'low': 'Past', 'medium': "O'rta", 'high': 'Yuqori'}
        elif language == 'uz_cyrl':
            title = "Тендер таҳлили ҳисоботи"
            tender_info_title = "Тендер маълумотлари"
            purpose_label = "Мақсад"
            type_label = "Тур"
            req_count_label = "Талаблар"
            ranking_title = "Иштирокчилар рейтинги"
            winner_title = "Ғолиб"
            score_label = "Балл"
            match_label = "Мослик"
            risk_label = "Хавф"
            strengths_label = "Кучли томонлар"
            weaknesses_label = "Камчиликлар"
            summary_title = "Хулоса"
            risk_levels = {'low': 'Паст', 'medium': "Ўрта", 'high': 'Юқори'}
        else:
            title = "Отчет анализа тендера"
            tender_info_title = "Информация о тендере"
            purpose_label = "Цель"
            type_label = "Тип"
            req_count_label = "Требования"
            ranking_title = "Рейтинг участников"
            winner_title = "Победитель"
            score_label = "Балл"
            match_label = "Соответствие"
            risk_label = "Риск"
            strengths_label = "Сильные стороны"
            weaknesses_label = "Слабые стороны"
            summary_title = "Заключение"
            risk_levels = {'low': 'Низкий', 'medium': 'Средний', 'high': 'Высокий'}
        
        # Sarlavha
        elements.append(Paragraph(title, title_style))
        
        # Dekorativ chiziq sarlavha ostida
        title_line = Table([['']], colWidths=[13*cm], style=TableStyle([
            ('LINEBELOW', (0, 0), (-1, 0), 1.5, PRIMARY),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ]))
        elements.append(title_line)
        elements.append(Spacer(1, 10))
        
        # G'olib - chiroyli ramkali quti
        if winner:
            winner_name = winner.get('participant_name', '-')
            winner_score = winner.get('total_weighted_score') or winner.get('overall_match_percentage', 0)
            
            # G'olib ikonkasi va nomi
            trophy_text = "🏆"
            winner_title_text = f"<b>{winner_title}</b>"
            winner_name_text = f"<font size='12'>{winner_name}</font>"
            winner_score_text = f"<font size='16' color='#336644'><b>{winner_score:.0f}%</b></font>"
            
            winner_content = [
                [Paragraph(winner_title_text, ParagraphStyle('WT', fontSize=9, textColor=GOLD, alignment=1, fontName=fonts['bold']))],
                [Paragraph(winner_name_text, ParagraphStyle('WN', fontSize=11, textColor=DARK, alignment=1, spaceBefore=2, fontName=fonts['bold']))],
                [Paragraph(winner_score_text, ParagraphStyle('WS', fontSize=12, alignment=1, spaceBefore=3, fontName=fonts['bold']))],
            ]
            
            winner_box = Table(winner_content, colWidths=[13*cm])
            winner_box.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), LIGHT_GREEN),
                ('BOX', (0, 0), (-1, -1), 1, PRIMARY),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('BOTTOMPADDING', (0, -1), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ]))
            elements.append(winner_box)
            elements.append(Spacer(1, 12))
        
        # Tender ma'lumotlari - chiroyli quti ichida
        if tender_analysis:
            # Bo'lim sarlavhasi
            section_header = Table(
                [[Paragraph(f"<b>{tender_info_title}</b>", ParagraphStyle('SH', fontSize=10, textColor=colors.white, fontName=fonts['bold']))]],
                colWidths=[13*cm]
            )
            section_header.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), PRIMARY),
                ('PADDING', (0, 0), (-1, -1), 5),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ]))
            elements.append(section_header)
            
            tender_purpose = tender_analysis.get('tender_purpose', '-')
            if len(tender_purpose) > 150:
                tender_purpose = tender_purpose[:150] + '...'
            
            # Ma'lumotlar jadvali
            info_data = [
                [Paragraph(f"<b>{purpose_label}:</b>", small_style), Paragraph(tender_purpose, normal_style)],
                [Paragraph(f"<b>{type_label}:</b>", small_style), Paragraph(str(tender_analysis.get('tender_type', '-')), normal_style)],
                [Paragraph(f"<b>{req_count_label}:</b>", small_style), 
                 Paragraph(f"{tender_analysis.get('requirements_count', 0)} (majburiy: {tender_analysis.get('mandatory_count', 0)})", normal_style)],
            ]
            
            info_table = Table(info_data, colWidths=[3*cm, 10*cm])
            info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), LIGHT_GRAY),
                ('BOX', (0, 0), (-1, -1), 0.5, BORDER),
                ('LINEABOVE', (0, 1), (-1, -1), 0.5, BORDER),
                ('PADDING', (0, 0), (-1, -1), 5),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            elements.append(info_table)
            elements.append(Spacer(1, 10))
        
        # Reyting jadvali
        if ranking:
            # Bo'lim sarlavhasi
            section_header2 = Table(
                [[Paragraph(f"<b>{ranking_title}</b>", ParagraphStyle('SH2', fontSize=10, textColor=colors.white, fontName=fonts['bold']))]],
                colWidths=[13*cm]
            )
            section_header2.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), PRIMARY),
                ('PADDING', (0, 0), (-1, -1), 5),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ]))
            elements.append(section_header2)
            
            # Jadval
            participant_col = 'Ishtirokchi' if language == 'uz_latn' else ('Иштирокчи' if language == 'uz_cyrl' else 'Участник')
            header = ['#', participant_col, score_label, match_label, risk_label]
            table_data = [header]
            
            for idx, p in enumerate(ranking, 1):
                score = p.get('total_weighted_score') or p.get('overall_match_percentage', 0)
                risk = risk_levels.get(p.get('risk_level', 'low'), p.get('risk_level', '-'))
                name = p.get('participant_name', '-')
                if len(name) > 30:
                    name = name[:30] + '...'
                
                row = [str(idx), name, f"{score:.0f}%", f"{p.get('overall_match_percentage', 0)}%", risk]
                table_data.append(row)
            
            ranking_table = Table(table_data, colWidths=[1*cm, 6.5*cm, 2*cm, 2*cm, 1.5*cm])
            
            table_style_list = [
                ('FONTNAME', (0, 0), (-1, 0), fonts['bold']),
                ('FONTNAME', (0, 1), (-1, -1), fonts['regular']),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BACKGROUND', (0, 0), (-1, 0), LIGHT_GRAY),
                ('TEXTCOLOR', (0, 0), (-1, 0), PRIMARY),
                ('TEXTCOLOR', (0, 1), (-1, -1), DARK),
                ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                ('ALIGN', (2, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('PADDING', (0, 0), (-1, -1), 6),
                ('BOX', (0, 0), (-1, -1), 0.5, BORDER),
                ('LINEBELOW', (0, 0), (-1, 0), 1, PRIMARY),
                ('LINEBELOW', (0, 1), (-1, -2), 0.3, BORDER),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.Color(0.98, 0.98, 0.98)]),
            ]
            
            # G'olib qatori - maxsus stil
            if len(table_data) > 1:
                table_style_list.append(('BACKGROUND', (0, 1), (-1, 1), LIGHT_GREEN))
                table_style_list.append(('TEXTCOLOR', (0, 1), (-1, 1), PRIMARY))
                table_style_list.append(('FONTNAME', (0, 1), (-1, 1), fonts['bold']))
            
            ranking_table.setStyle(TableStyle(table_style_list))
            elements.append(ranking_table)
            elements.append(Spacer(1, 12))
        
        # Har bir ishtirokchi haqida qisqacha - kartochka ko'rinishida
        if ranking:
            details_title = (
                "Ishtirokchilar tafsiloti" if language == 'uz_latn'
                else ("Иштирокчилар тафсилоти" if language == 'uz_cyrl' else "Детали участников")
            )
            section_header3 = Table(
                [[Paragraph(f"<b>{details_title}</b>", ParagraphStyle('SH3', fontSize=10, textColor=colors.white, fontName=fonts['bold']))]],
                colWidths=[13*cm]
            )
            section_header3.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), PRIMARY),
                ('PADDING', (0, 0), (-1, -1), 5),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ]))
            elements.append(section_header3)
            elements.append(Spacer(1, 5))
            
            for idx, p in enumerate(ranking, 1):
                name = p.get('participant_name', '-')
                score = p.get('total_weighted_score') or p.get('overall_match_percentage', 0)
                risk = risk_levels.get(p.get('risk_level', 'low'), p.get('risk_level', '-'))
                
                # Kuchli tomonlar
                strengths = p.get('strengths', [])
                s_text = ', '.join(strengths[:3]) if strengths else '-'
                if len(s_text) > 80:
                    s_text = s_text[:80] + '...'
                
                # Kamchiliklar
                weaknesses = p.get('weaknesses', [])
                w_text = ', '.join(weaknesses[:3]) if weaknesses else '-'
                if len(w_text) > 80:
                    w_text = w_text[:80] + '...'
                
                # Participant kartochkasi
                is_winner = idx == 1
                card_bg = LIGHT_GREEN if is_winner else colors.white
                card_border = PRIMARY if is_winner else BORDER
                
                card_data = [
                    [Paragraph(f"<b>{idx}. {name}</b>", ParagraphStyle('CN', fontSize=9, textColor=PRIMARY)), 
                     Paragraph(f"<b>{score:.0f}%</b>", ParagraphStyle('CS', fontSize=10, textColor=PRIMARY, alignment=2))],
                    [Paragraph(f"<font color='#336644'>+</font> {s_text}", ParagraphStyle('ST', fontSize=7, textColor=GRAY)), ''],
                    [Paragraph(f"<font color='#994444'>-</font> {w_text}", ParagraphStyle('WK', fontSize=7, textColor=GRAY)), ''],
                ]
                
                participant_card = Table(card_data, colWidths=[10*cm, 3*cm])
                participant_card.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), card_bg),
                    ('BOX', (0, 0), (-1, -1), 0.5, card_border),
                    ('PADDING', (0, 0), (-1, -1), 4),
                    ('SPAN', (0, 1), (1, 1)),
                    ('SPAN', (0, 2), (1, 2)),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING', (0, 0), (-1, 0), 5),
                    ('BOTTOMPADDING', (0, -1), (-1, -1), 5),
                ]))
                elements.append(participant_card)
                elements.append(Spacer(1, 4))
        
        # Xulosa - yangi sahifada
        if summary:
            elements.append(PageBreak())  # 2-sahifadan boshlash
            
            section_header4 = Table(
                [[Paragraph(f"<b>{summary_title}</b>", ParagraphStyle('SH4', fontSize=10, textColor=colors.white, fontName=fonts['bold']))]],
                colWidths=[13*cm]
            )
            section_header4.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), PRIMARY),
                ('PADDING', (0, 0), (-1, -1), 5),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ]))
            elements.append(section_header4)
            
            # Markdown tozalash
            clean_summary = summary.replace('**', '').replace('*', '').replace('#', '').strip()
            
            # Qisqartirish
            if len(clean_summary) > 2500:
                clean_summary = clean_summary[:2500] + '...'
            
            # Xulosa matni - satrlar orasini qisqartirish
            summary_paragraphs = []
            for line in clean_summary.split('\n'):
                line = line.strip()
                if line:
                    summary_paragraphs.append([Paragraph(line, ParagraphStyle('SUM', fontSize=9, textColor=DARK, leading=11, spaceAfter=1, fontName=fonts['regular']))])
            
            if summary_paragraphs:
                summary_table = Table(summary_paragraphs, colWidths=[13*cm])
                summary_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), LIGHT_GRAY),
                    ('BOX', (0, 0), (-1, -1), 0.5, BORDER),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                elements.append(summary_table)
        
        # PDF yaratish
        doc.build(elements, onFirstPage=lambda c, d: draw_header_footer(c, d, language),
                  onLaterPages=lambda c, d: draw_header_footer(c, d, language))
        
        # Response
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        
        filename = f"tender_tahlil_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        logger.error(f"PDF eksportda xatolik: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ============================================
# TAHLIL NATIJALARINI SAQLASH VA OLISH API
# ============================================

from .models import TenderAnalysisResult


@api_view(['POST'])
@permission_classes([AllowAny])
def save_analysis_result(request):
    """
    Tahlil natijasini bazaga saqlash
    
    POST /api/evaluations/save-result/
    
    Body:
        - tender: Tender tahlili (dict)
        - participants: Ishtirokchilar tahlili (list)
        - ranking: Reyting (list)
        - winner: G'olib (dict)
        - summary: Xulosa (str)
        - language: Til (str)
    """
    try:
        tender = request.data.get('tender', {})
        participants = request.data.get('participants', [])
        ranking = request.data.get('ranking', [])
        winner = request.data.get('winner', {})
        summary = request.data.get('summary', '')
        language = _normalize_language(request.data.get('language', 'uz_latn'))
        
        # Tender nomi
        tender_name = tender.get('purpose', '') or tender.get('name', '') or 'Nomsiz tender'
        tender_type = tender.get('type', '')
        
        # G'olib ma'lumotlari
        winner_name = winner.get('participant_name', '') if winner else ''
        winner_score = winner.get('total_weighted_score', 0) if winner else 0
        
        # Bazaga saqlash
        result = TenderAnalysisResult.objects.create(
            tender_name=tender_name,
            tender_type=tender_type,
            tender_data=tender,
            participants=participants,
            participant_count=len(participants),
            ranking=ranking,
            winner_name=winner_name,
            winner_score=winner_score,
            summary=summary,
            language=language
        )
        
        return Response({
            'success': True,
            'id': result.id,
            'message': 'Natija muvaffaqiyatli saqlandi'
        }, status=status.HTTP_201_CREATED)
        
    except Exception as e:
        logger.error(f"Natija saqlashda xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def get_analysis_history(request):
    """
    Tahlil tarixini olish
    
    GET /api/evaluations/history/
    
    Query params:
        - limit: Nechta natija (default: 20)
        - offset: Qayerdan boshlash (default: 0)
    """
    try:
        limit = int(request.GET.get('limit', 20))
        offset = int(request.GET.get('offset', 0))
        
        total = TenderAnalysisResult.objects.count()
        results = TenderAnalysisResult.objects.all()[offset:offset + limit]
        
        history = []
        for r in results:
            history.append({
                'id': r.id,
                'date': r.created_at.isoformat(),
                'tender': r.tender_name,
                'tender_type': r.tender_type,
                'winner': r.winner_name,
                'winner_score': r.winner_score,
                'participantCount': r.participant_count,
                'ranking': r.ranking,
                'summary': r.summary,
            })
        
        return Response({
            'success': True,
            'total': total,
            'history': history
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        logger.error(f"Tarix olishda xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def get_analysis_detail(request, pk):
    """
    Bitta tahlil natijasini olish
    
    GET /api/evaluations/history/<id>/
    """
    try:
        result = TenderAnalysisResult.objects.get(pk=pk)
        
        return Response({
            'success': True,
            'result': {
                'id': result.id,
                'date': result.created_at.isoformat(),
                'tender': result.tender_data,
                'tender_name': result.tender_name,
                'tender_type': result.tender_type,
                'participants': result.participants,
                'participantCount': result.participant_count,
                'ranking': result.ranking,
                'winner': result.winner_name,
                'winner_score': result.winner_score,
                'summary': result.summary,
                'language': result.language,
            }
        }, status=status.HTTP_200_OK)
        
    except TenderAnalysisResult.DoesNotExist:
        return Response({
            'success': False,
            'error': 'Natija topilmadi'
        }, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        logger.error(f"Natija olishda xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['DELETE'])
@api_view(['DELETE', 'POST'])
def delete_analysis_result(request, pk):
    """
    Tahlil natijasini o'chirish
    
    DELETE /api/evaluations/history/<id>/
    POST /api/evaluations/history/<id>/delete/
    """
    try:
        result = TenderAnalysisResult.objects.get(pk=pk)
        result.delete()
        
        return Response({
            'success': True,
            'message': 'Natija o\'chirildi'
        }, status=status.HTTP_200_OK)
        
    except TenderAnalysisResult.DoesNotExist:
        return Response({
            'success': False,
            'error': 'Natija topilmadi'
        }, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        logger.error(f"Natija o'chirishda xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def get_dashboard_stats(request):
    """
    Dashboard uchun statistikalar
    
    GET /api/evaluations/dashboard-stats/
    """
    try:
        from django.db.models import Sum, Avg
        
        total_analyses = TenderAnalysisResult.objects.count()
        total_participants = TenderAnalysisResult.objects.aggregate(
            total=Sum('participant_count')
        )['total'] or 0
        
        # Oxirgi 30 kun
        from datetime import timedelta
        from django.utils import timezone
        
        last_30_days = timezone.now() - timedelta(days=30)
        recent_analyses = TenderAnalysisResult.objects.filter(
            created_at__gte=last_30_days
        ).count()
        
        return Response({
            'success': True,
            'stats': {
                'total_tenders': total_analyses,
                'active_tenders': recent_analyses,
                'total_participants': total_participants,
                'total_evaluations': total_analyses,
            }
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        logger.error(f"Statistika olishda xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
def download_excel(request):
    """
    Tahlil natijasini Excel formatida yuklab olish
    
    POST /api/evaluations/download-excel/
    """
    import pandas as pd
    from io import BytesIO
    
    try:
        data = request.data
        tender = data.get('tender', {})
        ranking = data.get('ranking', [])
        summary = data.get('summary', '')
        language = _normalize_language(data.get('language', 'uz_latn'))
        
        if not ranking:
            return Response({
                'success': False,
                'error': 'Ranking ma\'lumotlari yo\'q'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Excel buffer
        output = BytesIO()
        
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            workbook = writer.book
            
            # Formatlar
            header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#4F46E5',
                'font_color': 'white',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })
            cell_format = workbook.add_format({
                'border': 1,
                'align': 'left',
                'valign': 'vcenter',
                'text_wrap': True
            })
            number_format = workbook.add_format({
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'num_format': '0.0'
            })
            winner_format = workbook.add_format({
                'bold': True,
                'bg_color': '#10B981',
                'font_color': 'white',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })
            
            # === SHEET 1: Reyting ===
            if language != 'ru':
                headers = ['O\'rin', 'Ishtirokchi', 'Umumiy ball', 'Moslik %', 'Narx', 'Xavf', 'Tavsiya']
            else:
                headers = ['Место', 'Участник', 'Общий балл', 'Соответствие %', 'Цена', 'Риск', 'Рекомендация']
            
            ranking_data = []
            for i, p in enumerate(ranking):
                ranking_data.append([
                    i + 1,
                    p.get('participant_name', ''),
                    p.get('total_weighted_score', 0),
                    p.get('overall_match_percentage', 0),
                    p.get('price_analysis', {}).get('proposed_price', 'N/A'),
                    p.get('risk_level', 'N/A'),
                    p.get('recommendation', '')[:100] + '...' if len(p.get('recommendation', '')) > 100 else p.get('recommendation', '')
                ])
            
            df_ranking = pd.DataFrame(ranking_data, columns=headers)
            df_ranking.to_excel(writer, sheet_name='Reyting' if language != 'ru' else 'Рейтинг', index=False, startrow=1)
            
            worksheet_ranking = writer.sheets['Reyting' if language != 'ru' else 'Рейтинг']
            
            # Title
            tender_name = tender.get('tender_purpose', 'Tender tahlili')[:50]
            worksheet_ranking.merge_range('A1:G1', tender_name, header_format)
            
            # Header formatting
            for col_num, value in enumerate(headers):
                worksheet_ranking.write(1, col_num, value, header_format)
            
            # Data formatting
            for row_num, row_data in enumerate(ranking_data):
                for col_num, cell_data in enumerate(row_data):
                    fmt = winner_format if row_num == 0 and col_num in [0, 1, 2] else (number_format if col_num in [2, 3] else cell_format)
                    worksheet_ranking.write(row_num + 2, col_num, cell_data, fmt)
            
            # Column widths
            worksheet_ranking.set_column('A:A', 8)
            worksheet_ranking.set_column('B:B', 25)
            worksheet_ranking.set_column('C:D', 15)
            worksheet_ranking.set_column('E:E', 20)
            worksheet_ranking.set_column('F:F', 12)
            worksheet_ranking.set_column('G:G', 40)
            
            # === SHEET 2: Batafsil ===
            detail_sheet = 'Batafsil' if language != 'ru' else 'Подробно'
            
            if language != 'ru':
                detail_headers = ['Ishtirokchi', 'Kuchli tomonlar', 'Zaif tomonlar']
            else:
                detail_headers = ['Участник', 'Сильные стороны', 'Слабые стороны']
            
            detail_data = []
            for p in ranking:
                strengths = '\n'.join(p.get('strengths', [])[:5])
                weaknesses = '\n'.join(p.get('weaknesses', [])[:5])
                detail_data.append([
                    p.get('participant_name', ''),
                    strengths,
                    weaknesses
                ])
            
            df_detail = pd.DataFrame(detail_data, columns=detail_headers)
            df_detail.to_excel(writer, sheet_name=detail_sheet, index=False)
            
            worksheet_detail = writer.sheets[detail_sheet]
            for col_num, value in enumerate(detail_headers):
                worksheet_detail.write(0, col_num, value, header_format)
            worksheet_detail.set_column('A:A', 25)
            worksheet_detail.set_column('B:C', 50)
            
            # === SHEET 3: Xulosa ===
            summary_sheet = 'Xulosa' if language != 'ru' else 'Итог'
            worksheet_summary = workbook.add_worksheet(summary_sheet)
            
            worksheet_summary.write('A1', 'Tahlil xulosasi' if language != 'ru' else 'Итог анализа', header_format)
            worksheet_summary.merge_range('A2:D20', summary, cell_format)
            worksheet_summary.set_column('A:D', 30)
        
        output.seek(0)
        
        # Response
        filename = f"tender_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        response = HttpResponse(
            output.read(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        # Audit log (agar user bo'lsa)
        try:
            from apps.users.models import AuditLog
            if request.user.is_authenticated:
                AuditLog.log(
                    request.user,
                    AuditLog.ActionType.EXCEL_DOWNLOAD,
                    f'Excel yuklab olindi: {tender_name}',
                    request
                )
        except:
            pass
        
        return response
        
    except Exception as e:
        logger.error(f"Excel yaratishda xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([AllowAny])
def download_csv(request):
    """
    Tahlil natijasini CSV formatida yuklab olish
    
    POST /api/evaluations/download-csv/
    """
    import csv
    from io import StringIO
    
    try:
        data = request.data
        ranking = data.get('ranking', [])
        language = _normalize_language(data.get('language', 'uz_latn'))
        
        if not ranking:
            return Response({
                'success': False,
                'error': 'Ranking ma\'lumotlari yo\'q'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # CSV buffer
        output = StringIO()
        
        if language != 'ru':
            headers = ['O\'rin', 'Ishtirokchi', 'Umumiy ball', 'Moslik %', 'Narx', 'Xavf', 'Tavsiya']
        else:
            headers = ['Место', 'Участник', 'Общий балл', 'Соответствие %', 'Цена', 'Риск', 'Рекомендация']
        
        writer = csv.writer(output)
        writer.writerow(headers)
        
        for i, p in enumerate(ranking):
            writer.writerow([
                i + 1,
                p.get('participant_name', ''),
                p.get('total_weighted_score', 0),
                p.get('overall_match_percentage', 0),
                p.get('price_analysis', {}).get('proposed_price', 'N/A'),
                p.get('risk_level', 'N/A'),
                p.get('recommendation', '')
            ])
        
        output.seek(0)
        
        filename = f"tender_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        response = HttpResponse(output.read(), content_type='text/csv; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        logger.error(f"CSV yaratishda xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def get_chart_data(request):
    """
    Dashboard grafiklar uchun ma'lumotlar
    
    GET /api/evaluations/chart-data/
    """
    try:
        from django.db.models import Count, Avg
        from django.db.models.functions import TruncDate, TruncMonth
        from datetime import timedelta
        from django.utils import timezone
        
        # Oxirgi 30 kun
        last_30_days = timezone.now() - timedelta(days=30)
        last_12_months = timezone.now() - timedelta(days=365)
        
        # Kunlik tahlillar (oxirgi 30 kun)
        daily_analyses = TenderAnalysisResult.objects.filter(
            created_at__gte=last_30_days
        ).annotate(
            date=TruncDate('created_at')
        ).values('date').annotate(
            count=Count('id'),
            avg_score=Avg('winner_score')
        ).order_by('date')
        
        # Oylik tahlillar (oxirgi 12 oy)
        monthly_analyses = TenderAnalysisResult.objects.filter(
            created_at__gte=last_12_months
        ).annotate(
            month=TruncMonth('created_at')
        ).values('month').annotate(
            count=Count('id'),
            total_participants=Count('participant_count')
        ).order_by('month')
        
        # Tender turlari bo'yicha
        tender_types = TenderAnalysisResult.objects.values(
            'tender_type'
        ).annotate(
            count=Count('id')
        ).order_by('-count')[:5]
        
        # G'oliblar ball taqsimoti
        score_distribution = []
        score_ranges = [(0, 50), (50, 60), (60, 70), (70, 80), (80, 90), (90, 100)]
        for low, high in score_ranges:
            count = TenderAnalysisResult.objects.filter(
                winner_score__gte=low,
                winner_score__lt=high
            ).count()
            score_distribution.append({
                'range': f'{low}-{high}',
                'count': count
            })
        
        # Ishtirokchilar soni taqsimoti
        participant_distribution = TenderAnalysisResult.objects.values(
            'participant_count'
        ).annotate(
            count=Count('id')
        ).order_by('participant_count')[:10]
        
        return Response({
            'success': True,
            'charts': {
                'daily_analyses': list(daily_analyses),
                'monthly_analyses': list(monthly_analyses),
                'tender_types': list(tender_types),
                'score_distribution': score_distribution,
                'participant_distribution': list(participant_distribution)
            }
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        logger.error(f"Grafik ma'lumotlari olishda xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def get_recent_activities(request):
    """
    So'nggi faoliyatlar (Audit log)
    
    GET /api/evaluations/recent-activities/
    """
    try:
        activities = []
        
        # So'nggi tahlillar
        recent_analyses = TenderAnalysisResult.objects.order_by('-created_at')[:10]
        for analysis in recent_analyses:
            activities.append({
                'type': 'analysis',
                'title': analysis.tender_name[:50] if analysis.tender_name else 'Tender tahlili',
                'description': f"G'olib: {analysis.winner_name} ({analysis.winner_score:.1f} ball)",
                'date': analysis.created_at.isoformat(),
                'participant_count': analysis.participant_count
            })
        
        # Audit loglar (agar mavjud bo'lsa)
        try:
            from apps.users.models import AuditLog
            recent_logs = AuditLog.objects.order_by('-created_at')[:10]
            for log in recent_logs:
                activities.append({
                    'type': 'audit',
                    'title': log.get_action_display(),
                    'description': log.description,
                    'date': log.created_at.isoformat(),
                    'user': log.user.username if log.user else 'System'
                })
        except:
            pass
        
        # Vaqt bo'yicha tartiblash
        activities.sort(key=lambda x: x['date'], reverse=True)
        
        return Response({
            'success': True,
            'activities': activities[:20]
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        logger.error(f"Faoliyatlar olishda xatolik: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
