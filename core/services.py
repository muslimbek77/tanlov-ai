import os
import logging
import magic
from typing import Dict, List, Any, Optional
from pathlib import Path
import PyPDF2
import pandas as pd
from unstructured.partition.auto import partition
from unstructured.documents.elements import Element
import openpyxl
from PIL import Image
import pytesseract
from datetime import datetime
import hashlib

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Smart Ingest Moduli - Hujjatlarni qayta ishlash xizmati"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'xlsx', 'xls', 'docx', 'txt', 'jpg', 'jpeg', 'png', 'tiff']
    
    def process_document(self, file_path: str, document_type: str = 'auto') -> Dict[str, Any]:
        """
        Hujjatni qayta ishlash
        """
        try:
            file_path = Path(file_path)
            
            # Fayl mavjudligini tekshirish
            if not file_path.exists():
                raise FileNotFoundError(f"Fayl topilmadi: {file_path}")
            
            # Fayl turini aniqlash
            file_type = self._detect_file_type(file_path)
            
            if file_type not in self.supported_formats:
                raise ValueError(f"Qo'llab-quvvatlanmaydigan fayl turi: {file_type}")
            
            # Metadata olish
            metadata = self._extract_metadata(file_path, file_type)
            
            # Matnni ajratish
            text_content = self._extract_text(file_path, file_type)
            
            # Strukturalash
            structured_data = self._structure_content(text_content, document_type)
            
            # Vektorlashtirish uchun ma'lumotlar
            vector_data = self._prepare_vector_data(structured_data, metadata)
            
            result = {
                'file_path': str(file_path),
                'file_type': file_type,
                'metadata': metadata,
                'text_content': text_content,
                'structured_data': structured_data,
                'vector_data': vector_data,
                'processing_status': 'success',
                'processed_at': datetime.now().isoformat(),
            }
            
            logger.info(f"Hujjat muvaffaqiyatli qayta ishlindi: {file_path}")
            return result
            
        except Exception as e:
            logger.error(f"Hujjatni qayta ishlashda xatolik: {str(e)}")
            return {
                'file_path': str(file_path),
                'processing_status': 'error',
                'error': str(e),
                'processed_at': datetime.now().isoformat(),
            }
    
    def _detect_file_type(self, file_path: Path) -> str:
        """Fayl turini aniqlash"""
        mime_type = magic.from_file(str(file_path), mime=True)
        
        if mime_type == 'application/pdf':
            return 'pdf'
        elif mime_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
            return 'xlsx'
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return 'docx'
        elif mime_type.startswith('text/'):
            return 'txt'
        elif mime_type.startswith('image/'):
            extension = file_path.suffix.lower().lstrip('.')
            return extension if extension in ['jpg', 'jpeg', 'png', 'tiff'] else 'jpg'
        
        # Fallback to extension
        return file_path.suffix.lower().lstrip('.')
    
    def _extract_metadata(self, file_path: Path, file_type: str) -> Dict[str, Any]:
        """Metadata ajratish"""
        metadata = {
            'file_name': file_path.name,
            'file_size': file_path.stat().st_size,
            'file_type': file_type,
            'creation_date': datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            'modification_date': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
        }
        
        try:
            if file_type == 'pdf':
                metadata.update(self._extract_pdf_metadata(file_path))
            elif file_type in ['xlsx', 'xls']:
                metadata.update(self._extract_excel_metadata(file_path))
            elif file_type == 'docx':
                metadata.update(self._extract_docx_metadata(file_path))
        except Exception as e:
            logger.warning(f"Metadata ajratishda xatolik: {str(e)}")
            metadata['metadata_error'] = str(e)
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """PDF metadata ajratish"""
        metadata = {}
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata.update({
                    'page_count': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                    'creator': pdf_reader.metadata.get('/Creator', '') if pdf_reader.metadata else '',
                    'producer': pdf_reader.metadata.get('/Producer', '') if pdf_reader.metadata else '',
                    'creation_date': pdf_reader.metadata.get('/CreationDate', '') if pdf_reader.metadata else '',
                    'modification_date': pdf_reader.metadata.get('/ModDate', '') if pdf_reader.metadata else '',
                })
                
                # PDF versiyasi
                metadata['pdf_version'] = pdf_reader.pdf_header
                
        except Exception as e:
            logger.error(f"PDF metadata ajratishda xatolik: {str(e)}")
        
        return metadata
    
    def _extract_excel_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Excel metadata ajratish"""
        metadata = {}
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True)
            
            metadata.update({
                'sheet_count': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames,
                'max_row': 0,
                'max_column': 0,
            })
            
            # Har bir varaq uchun ma'lumot
            sheet_info = {}
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_info[sheet_name] = {
                    'max_row': sheet.max_row,
                    'max_column': sheet.max_column,
                    'has_data': sheet.max_row > 1 and sheet.max_column > 1,
                }
                
                if sheet.max_row > metadata['max_row']:
                    metadata['max_row'] = sheet.max_row
                if sheet.max_column > metadata['max_column']:
                    metadata['max_column'] = sheet.max_column
            
            metadata['sheet_info'] = sheet_info
            workbook.close()
            
        except Exception as e:
            logger.error(f"Excel metadata ajratishda xatolik: {str(e)}")
        
        return metadata
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """DOCX metadata ajratish"""
        metadata = {}
        try:
            from docx import Document
            doc = Document(file_path)
            
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
            })
            
        except Exception as e:
            logger.error(f"DOCX metadata ajratishda xatolik: {str(e)}")
        
        return metadata
    
    def _extract_text(self, file_path: Path, file_type: str) -> str:
        """Matnni ajratish"""
        try:
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type in ['xlsx', 'xls']:
                return self._extract_excel_text(file_path)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path)
            elif file_type == 'txt':
                return self._extract_text_file(file_path)
            elif file_type in ['jpg', 'jpeg', 'png', 'tiff']:
                return self._extract_image_text(file_path)
            else:
                # Unstructured.io dan foydalanish
                elements = partition(filename=str(file_path))
                return '\n'.join([str(element) for element in elements])
        
        except Exception as e:
            logger.error(f"Matn ajratishda xatolik: {str(e)}")
            return ""
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """PDF dan matn ajratish"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"PDF matn ajratishda xatolik: {str(e)}")
        
        return text
    
    def _extract_excel_text(self, file_path: Path) -> str:
        """Excel dan matn ajratish"""
        text = ""
        try:
            # Unstructured.io dan foydalanish
            elements = partition(filename=str(file_path))
            text = '\n'.join([str(element) for element in elements])
            
            # Agar unstructured ishlamasa
            if not text.strip():
                df = pd.read_excel(file_path, sheet_name=None)
                for sheet_name, sheet_df in df.items():
                    text += f"\n\n=== {sheet_name} ===\n"
                    text += sheet_df.to_string()
        
        except Exception as e:
            logger.error(f"Excel matn ajratishda xatolik: {str(e)}")
        
        return text
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """DOCX dan matn ajratish"""
        text = ""
        try:
            from docx import Document
            doc = Document(file_path)
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text += " | ".join(row_text) + "\n"
        
        except Exception as e:
            logger.error(f"DOCX matn ajratishda xatolik: {str(e)}")
        
        return text
    
    def _extract_text_file(self, file_path: Path) -> str:
        """Text fayldan matn olish"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Boshqa kodlashlarni sinab ko'rish
            for encoding in ['latin-1', 'cp1252', 'cp1251']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            return ""
    
    def _extract_image_text(self, file_path: Path) -> str:
        """Rasmdan matn ajratish (OCR)"""
        text = ""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='uzb+rus+eng')
        except Exception as e:
            logger.error(f"OCR da xatolik: {str(e)}")
        
        return text
    
    def _structure_content(self, text_content: str, document_type: str) -> Dict[str, Any]:
        """Kontentni strukturalash"""
        structured = {
            'document_type': document_type,
            'sections': [],
            'tables': [],
            'key_information': {},
            'entities': [],
        }
        
        try:
            # Unstructured.io elementlari
            if text_content:
                # Bo'limlarga ajratish
                lines = text_content.split('\n')
                current_section = None
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Sarlavhalarni aniqlash
                    if self._is_heading(line):
                        if current_section:
                            structured['sections'].append(current_section)
                        current_section = {
                            'title': line,
                            'content': [],
                            'level': self._get_heading_level(line),
                        }
                    elif current_section:
                        current_section['content'].append(line)
                    else:
                        if not structured['sections']:
                            structured['sections'].append({
                                'title': 'Kirish',
                                'content': [line],
                                'level': 1,
                            })
                        else:
                            structured['sections'][0]['content'].append(line)
                
                if current_section:
                    structured['sections'].append(current_section)
                
                # Kalit so'zlarni ajratish
                structured['key_information'] = self._extract_key_information(text_content)
                
        except Exception as e:
            logger.error(f"Strukturalashda xatolik: {str(e)}")
        
        return structured
    
    def _is_heading(self, line: str) -> bool:
        """Sarlavha ekanligini tekshirish"""
        # Oddiy qoidalar
        if len(line) < 100 and (
            line.isupper() or
            line.endswith(':') or
            any(keyword in line.lower() for keyword in [
                'tender', 'taklif', 'shartnoma', 'talab', 'kiritish', 'xulosa',
                'тендер', 'предложение', 'договор', 'требование', 'введение', 'заключение'
            ])
        ):
            return True
        return False
    
    def _get_heading_level(self, line: str) -> int:
        """Sarlavha darajasini aniqlash"""
        if line.isupper():
            return 1
        elif line.startswith(' '):
            return 3
        else:
            return 2
    
    def _extract_key_information(self, text: str) -> Dict[str, Any]:
        """Kalit ma'lumotlarni ajratish"""
        info = {}
        
        # Raqamlarni ajratish
        import re
        
        # Summalarni topish
        amounts = re.findall(r'[\d,]+(?:\.\d{2})?\s?(?:UZS|USD|sum|dollar)', text, re.IGNORECASE)
        if amounts:
            info['amounts'] = amounts
        
        # Sanalarni topish
        dates = re.findall(r'\d{2}\.\d{2}\.\d{4}|\d{4}-\d{2}-\d{2}', text)
        if dates:
            info['dates'] = dates
        
        # Telefon raqamlarini topish
        phones = re.findall(r'\+?\d{1,3}?\s?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}', text)
        if phones:
            info['phone_numbers'] = phones
        
        # Email larni topish
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if emails:
            info['emails'] = emails
        
        return info
    
    def _prepare_vector_data(self, structured_data: Dict[str, Any], metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Vektorlashtirish uchun ma'lumotlarni tayyorlash"""
        vector_data = {
            'content_chunks': [],
            'metadata_chunks': [],
        }
        
        try:
            # Matn bo'laklariga ajratish
            all_text = ""
            for section in structured_data.get('sections', []):
                section_text = f"{section['title']}\n{' '.join(section['content'])}"
                all_text += section_text + "\n\n"
            
            # Bo'laklarga ajratish (chunking)
            chunk_size = 1000  # har bir bo'lak 1000 ta belgi
            chunks = [all_text[i:i+chunk_size] for i in range(0, len(all_text), chunk_size)]
            vector_data['content_chunks'] = chunks
            
            # Metadatani qo'shish
            metadata_text = f"File: {metadata.get('file_name', '')}\n"
            metadata_text += f"Type: {metadata.get('file_type', '')}\n"
            metadata_text += f"Size: {metadata.get('file_size', 0)} bytes\n"
            metadata_text += f"Created: {metadata.get('creation_date', '')}\n"
            
            if metadata.get('author'):
                metadata_text += f"Author: {metadata['author']}\n"
            if metadata.get('creator'):
                metadata_text += f"Creator: {metadata['creator']}\n"
            
            vector_data['metadata_chunks'] = [metadata_text]
            
        except Exception as e:
            logger.error(f"Vektor ma'lumotlarini tayyorlashda xatolik: {str(e)}")
        
        return vector_data


class VectorEmbeddingService:
    """Vektor embedding xizmati"""
    
    def __init__(self):
        self.embedding_model = None
        self._initialize_model()
    
    def _initialize_model(self):
        """Embedding modelini ishga tushirish"""
        try:
            # OpenAI embedding modeli
            from openai import OpenAI
            import os
            
            if os.getenv('OPENAI_API_KEY'):
                self.client = OpenAI()
                self.model_type = 'openai'
            else:
                # Local model (Sentence Transformers)
                from sentence_transformers import SentenceTransformer
                self.model = SentenceTransformer('all-MiniLM-L6-v2')
                self.model_type = 'local'
        
        except Exception as e:
            logger.error(f"Embedding modelini ishga tushirishda xatolik: {str(e)}")
            self.model_type = None
    
    def create_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Matnlar uchun embeddinglar yaratish"""
        if not self.model_type:
            raise ValueError("Embedding modeli mavjud emas")
        
        try:
            if self.model_type == 'openai':
                return self._create_openai_embeddings(texts)
            else:
                return self._create_local_embeddings(texts)
        
        except Exception as e:
            logger.error(f"Embedding yaratishda xatolik: {str(e)}")
            raise
    
    def _create_openai_embeddings(self, texts: List[str]) -> List[List[float]]:
        """OpenAI embeddinglari"""
        embeddings = []
        for text in texts:
            response = self.client.embeddings.create(
                model="text-embedding-3-small",
                input=text
            )
            embeddings.append(response.data[0].embedding)
        return embeddings
    
    def _create_local_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Local embeddinglar"""
        return self.model.encode(texts).tolist()


# Global xizmatlar
document_processor = DocumentProcessor()
embedding_service = VectorEmbeddingService()
